"""
Text extraction service.

Extracts raw text from uploaded PDF and DOCX documents,
cleans it, and returns a normalised string.
"""

from __future__ import annotations

import re
from pathlib import Path

import pdfplumber
from docx import Document


# ---------------------------------------------------------------------------
# Custom exception
# ---------------------------------------------------------------------------

class ExtractionError(Exception):
    """Raised when text extraction from a document fails."""

    def __init__(self, detail: str, source: str = "") -> None:
        self.detail = detail
        self.source = source
        super().__init__(detail)


# ---------------------------------------------------------------------------
# Public API
# ---------------------------------------------------------------------------

def extract_text(file_path: str) -> str:
    """
    Extract and clean text from *file_path*.

    Supported formats:
    - **.pdf** — via ``pdfplumber``
    - **.docx** — via ``python-docx``

    Returns the cleaned, normalised text.

    Raises
    ------
    ExtractionError
        If the file type is unsupported or extraction fails.
    """

    path = Path(file_path)

    if not path.exists():
        raise ExtractionError(
            f"File not found: {file_path}",
            source=file_path,
        )

    ext = path.suffix.lower()

    try:
        if ext == ".pdf":
            raw = _extract_pdf(path)
        elif ext == ".docx":
            raw = _extract_docx(path)
        else:
            raise ExtractionError(
                f"Unsupported file extension '{ext}' for text extraction. "
                "Only .pdf and .docx are supported.",
                source=file_path,
            )
    except ExtractionError:
        raise
    except Exception as exc:
        raise ExtractionError(
            f"Failed to extract text from '{path.name}': {exc}",
            source=file_path,
        ) from exc

    cleaned = _clean_text(raw)

    if not cleaned:
        raise ExtractionError(
            f"No readable text could be extracted from '{path.name}'.",
            source=file_path,
        )

    return cleaned


# ---------------------------------------------------------------------------
# Format-specific extractors
# ---------------------------------------------------------------------------

def _extract_pdf(path: Path) -> str:
    """Extract text from every page of a PDF using pdfplumber."""

    pages_text: list[str] = []

    with pdfplumber.open(path) as pdf:
        for page_num, page in enumerate(pdf.pages, start=1):
            text = page.extract_text()
            if text:
                pages_text.append(text)

    return "\n\n".join(pages_text)


def _extract_docx(path: Path) -> str:
    """Extract text from all paragraphs and tables of a DOCX file."""

    doc = Document(str(path))
    parts: list[str] = []

    # Paragraphs
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            parts.append(text)

    # Tables — flatten each cell
    for table in doc.tables:
        for row in table.rows:
            row_cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if row_cells:
                parts.append(" | ".join(row_cells))

    return "\n\n".join(parts)


# ---------------------------------------------------------------------------
# Text cleaning
# ---------------------------------------------------------------------------

def _clean_text(text: str) -> str:
    """
    Normalise extracted text:
    - Collapse multiple spaces / tabs into a single space.
    - Normalise line breaks (3+ consecutive newlines → 2).
    - Strip leading / trailing whitespace per line.
    - Strip leading / trailing whitespace from the whole string.
    """

    # Replace tabs with spaces
    text = text.replace("\t", " ")

    # Collapse multiple spaces into one
    text = re.sub(r"[^\S\n]+", " ", text)

    # Strip each line individually
    lines = [line.strip() for line in text.splitlines()]
    text = "\n".join(lines)

    # Collapse 3+ consecutive newlines into exactly 2
    text = re.sub(r"\n{3,}", "\n\n", text)

    return text.strip()
